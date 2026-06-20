#!/usr/bin/env python3
"""Align SOCVault-Investor-Partner-Deck.docx with blueprint ADR-006 (June 2026). Idempotent through v1.16."""
from docx import Document
from docx.text.paragraph import Paragraph
from copy import deepcopy

DECK = "/Users/samihaider/Documents/GitHub/SOCVault-Blueprint/BusinessProposal/SOCVault-Investor-Partner-Deck.docx"

REPLACEMENTS = [
    ("Financial model: 99.6% gross margin confirmed", "Financial model: 97.6% VAPT gross margin confirmed ($0.36 COGS)"),
    ("Estimated AI COGS per VAPT scan with caching: $0.052 on $15.00 revenue = 99.6% gross margin",
     "Fully-loaded COGS per VAPT scan (Claude + compute): ~$0.36 on $15.00 revenue = 97.6% gross margin"),
    ("Gross margin per scan: 99.6%", "Gross margin per VAPT scan: 97.6%"),
    ("AI cost deflation: Claude API pricing with prompt caching makes per-scan AI COGS ~$0.052 — economics that didn't exist 24 months ago",
     "AI cost deflation: Claude API with prompt caching keeps fully-loaded VAPT COGS at ~$0.36 — 97.6% gross margin at $15/scan"),
    ("2.  AI cost efficiency: Claude API with prompt caching makes per-scan COGS $0.052 — viable at $15/scan with 99.6% gross margin.",
     "2.  AI cost efficiency: Fully-loaded VAPT COGS is ~$0.36 (Claude + compute) — 97.6% gross margin at $15/scan."),
    ("SOCVault's estimated AWS monthly cost at launch: $800–$1,200/month. At $100K credits, infrastructure costs are covered through the first 7–8 months of live customer operations.",
     "SOCVault's estimated AWS monthly cost at launch: ~$100–$115/month (serverless staging MVP on Free Tier). Paid tier (CDN, WAF, EKS) scales from ~$433/month. $100K Activate credits cover 12+ months of early ops."),
    ("Business email + phone OTP → first scan result in under 5 minutes. No credit card. No DNS verification. No agent installation. No sales call. Every competitor requires 15 minutes to several weeks.",
     "Business email + phone OTP → first L1 scan result in under 5 minutes. No credit card for freemium. Domain ownership verified before paid scans. No agent installation for L1–L6. No sales call."),
    ("USP 4 — Full 8-Layer Coverage", "USP 4 — Full 8-Layer Core Coverage (+ L9 AI Agent Scan)"),
    ("Every SMB competitor covers 1–2 attack surface layers. SOCVault covers all 8 in one platform. Equivalent multi-vendor stack costs $34,920/year vs SOCVault at $2,388/year.",
     "Every SMB competitor covers 1–2 attack surface layers. SOCVault covers all 8 core layers in one platform, with L9 autonomous AI agent scanning in Phase 2. Equivalent multi-vendor stack costs $34,920/year vs SOCVault at $2,388/year."),
    ("SOCVault wraps eight scanning layers behind a single multi-tenant SaaS interface.",
     "SOCVault wraps eight core scanning layers (plus L9 AI Agent Scan in Phase 2) behind a single multi-tenant SaaS interface."),
    ("One platform. Eight security layers. AI-enabled detection, response, and financial risk translation.",
     "One platform. Eight core security layers + L9 AI Agent Scan. AI-enabled detection, response, and financial risk translation."),
    ("SOCVault is the unified AI-enabled cybersecurity solution for SMBs — one platform covering all eight attack surface layers at a price point every small business can afford.",
     "SOCVault is the unified AI-enabled cybersecurity solution for SMBs — one platform covering eight core attack surface layers plus L9 AI Agent Scan at a price point every small business can afford."),
    # Architecture — catch legacy and prior-deck strings
    ("Each scan executes in an ephemeral AWS Fargate task — no shared memory or persistent state between tenants",
     "L1 on AWS Lambda; L2+ on isolated Fargate/EKS jobs; API on API Gateway + Lambda — no shared memory between tenants"),
    ("L1 runs in AWS Lambda; L2–L6 run in isolated Fargate scan tasks; API/Celery on auto-scaling ECS Fargate — no shared memory between tenants",
     "L1 on AWS Lambda; L2+ on isolated scan workers; MVP API on API Gateway + Lambda + SQS; EKS at paid tier — tenant isolation by design"),
    ("Celery + Redis", "SQS + Lambda"),
    ("Celery / Redis", "SQS + Lambda"),
    ("L1 runs in AWS Lambda; L2–L6 run in isolated scan workers (Fargate at scale; Free Tier EC2 Docker in MVP) — no shared memory between tenants",
     "L1 on AWS Lambda; L2+ on isolated scan workers; MVP API on API Gateway + Lambda + SQS; EKS at paid tier — tenant isolation by design"),
    ("All secrets in AWS Secrets Manager — zero credentials in source code or environment files",
     "Secrets in AWS SSM Parameter Store (MVP) → Secrets Manager at paid tier — zero credentials in source code"),
    ("Local MVP — scanners + AI + dashboard running in Docker", "AWS staging MVP — API Gateway + Lambda + Amplify on api-staging.socvault.io"),
    ("AWS staging MVP — CI/CD + auth + L1 + Claude on api-staging.socvault.io",
     "AWS staging MVP — API Gateway + Lambda + Amplify; automated QA; production dormant until cutover"),
    ("Local MVP · 50 paying clients · $10K MRR", "AWS staging MVP · 50 paying clients · $10K MRR"),
    ("AWS deployment live", "Production cutover live"),
    ("api.socvault.io live; Fargate scans running; Cognito auth working",
     "api-staging green; API Gateway + Lambda + Cognito; OpenAPI + 208 user stories; Terraform IaC"),
    ("api.socvault.io live; ECS Fargate auto-scaling; Cognito auth; OpenAPI contract published",
     "api-staging green; API Gateway + Lambda + Amplify; Cognito; automated QA on every deploy"),
    ("api.socvault.io live; dual-env on single EC2; Cognito auth; OpenAPI contract published",
     "api-staging green; serverless stack; production dormant until cutover checklist"),
    ("COGS per Web VAPT scan (with Claude prompt caching)", "COGS per Web VAPT scan (Claude + compute, prompt caching)"),
    ("COGS per Web VAPT scan (Claude + Fargate, prompt caching)", "COGS per Web VAPT scan (Claude + compute, prompt caching)"),
    ("Multi-tenant REST API · JWT auth · Pydantic validation", "Multi-tenant REST API · Cognito JWT · API Gateway · Pydantic"),
    ("AWS eu-west-2 · ALB · S3 · SNS · Secrets Manager · CloudFront",
     "AWS eu-west-2 · API Gateway · Lambda · Amplify · SQS · SSM · S3 · Cognito · Atlas M0"),
    ("AWS eu-west-2 · EC2 (Free Tier) · Lambda · SSM · S3 · SNS · Cognito · staging + production",
     "AWS eu-west-2 · API Gateway · Lambda · Amplify · SQS · SSM · S3 · Cognito · staging MVP"),
    ("AWS eu-west-2 · ALB · ECS Fargate · Lambda · SSM · S3 · SNS · Cognito · staging + production",
     "AWS eu-west-2 · API Gateway · Lambda · Amplify · SQS · SSM · S3 · Cognito · staging MVP"),
    ("Terraform · GitHub Actions → ECR → ECS", "Terraform · GitHub Actions → Lambda/Amplify staging · CodePipeline (paid)"),
    ("GitHub Actions → ECR → EC2 Docker (staging + production) · Terraform (Phase 2b+)",
     "GitHub Actions → Terraform + Lambda/Amplify (staging) · CodePipeline (paid tier)"),
    ("GitHub Actions → ECR → ECS Fargate rolling deploy (staging + production) · Terraform",
     "GitHub Actions → Terraform + serverless staging deploy · EKS + CodePipeline (paid tier)"),
    ("$14.95", "$14.64"),
    ("L2–L6 VAPT analysis", "L2–L6 VAPT (Claude + compute)"),
    ("L2–L6 VAPT (Claude + Fargate)", "L2–L6 VAPT (Claude + compute)"),
    ("SOCVault SOC Pro (all 8 layers)", "SOCVault SOC Pro (all 8 core layers + L9)"),
    ("8/8 ✅", "8+L9 ✅"),
    ("$25,000 || 6 months Fargate ECS scan runtime + Lambda + S3 + Cognito",
     "$25,000 || 12+ months serverless MVP (API GW + Lambda + Amplify) + S3 + Cognito"),
    ("$25,000 || 12+ months Free Tier EC2 + Lambda + S3 + Cognito + staging/production",
     "$25,000 || 12+ months serverless MVP + Lambda + S3 + Cognito + staging"),
    ("$25,000 || 12+ months ECS Fargate auto-scaling + Lambda + S3 + Cognito + staging/production",
     "$25,000 || 12+ months serverless MVP + Lambda + S3 + Cognito + staging"),
    ("MobSF binary analysis · MASVS report · AI risk summary || 97.6%",
     "MobSF binary analysis · MASVS report · AI risk summary || 97.8%"),
    ("CloudFox/Pacu audit · IAM risk map · Privilege escalation paths || 97.6%",
     "CloudFox/Pacu audit · IAM risk map · Privilege escalation paths || 97.8%"),
    ("99.6%", "97.6%"),
    ("Cloudflare WAF · DDoS · CDN", "CloudFront · AWS WAF · CDN (paid tier)"),
    ("Multi-tenant user pools with custom tenant_id JWT attribute + SMS OTP via SNS",
     "One Cognito User Pool per environment; tenant_id custom attribute + SMS OTP via SNS"),
    ("per-tenant user pools", "one Cognito User Pool per environment (tenant_id attribute)"),
    ("Production promotion live", "Production cutover live"),
    ("September 2026 (Week 12)", "Week 17+ (cutover checklist)"),
    ("dual-env single account", "single account · staging active · production dormant"),
    ("200 clients · $30K MRR · AWS Marketplace live", "~200 clients · ~$16K MRR at Seed close · AWS Marketplace live"),
    ("200 clients · $30K MRR", "~200 clients · ~$16K MRR at Seed close"),
    # v1.14 — deck audit (Seed deck, milestones, blueprint honesty)
    ("Series A deck initiated", "Seed deck initiated"),
    ("Series A deck ready at Month 15", "Seed deck ready at Month 15"),
    ("Series A deck ready", "Seed deck ready"),
    ("12 months runway to $10K MRR and Series A positioning", "12 months runway to $10K MRR and Seed round positioning"),
    ("12-month runway · 50 paying clients · $10K MRR · Series A positioning",
     "12-month runway · 50 paying clients · $10K MRR · Seed round positioning"),
    ("deliver a working product to 50 paying clients",
     "deliver staging MVP to 50 paying clients (blueprint complete; socvault-app build next)"),
    ("Fully-loaded COGS per VAPT scan (Claude + Fargate):",
     "Fully-loaded COGS per VAPT scan (Claude + compute):"),
    ("see blueprint doc 20 for TI catalogue", "see `20_FREE_EXTERNAL_APIS.md` for Threat Intel catalogue"),
    ("Financial model: 97.6% VAPT gross margin confirmed ($0.36 COGS), break-even at Month 28, LTV:CAC = 12:1",
     "Financial model: 97.6% VAPT gross margin confirmed ($0.36 COGS), EBITDA ~Month 18 · cash-flow break-even Month 28, LTV:CAC = 12:1"),
    ("50 paying clients | December 2026 (Week 24) | $10K MRR milestone; SOAR live with real Wazuh alerts",
     "50 paying clients · $10K MRR | February 2027 (~M8) | Seed unlock criteria; SOAR live with real Wazuh alerts"),
    ("$10K MRR | Q1 2027 | Product-market fit validation; Seed deck initiated",
     "Seed unlock (included above) | February 2027 (~M8) | 50 clients + $10K MRR; Seed deck ready"),
    ("$50K MRR / Series A ready | Q3 2027 | 200+ clients; 3 MSP partners; Seed round close",
     "323 paying clients · $22K MRR (M18) | November 2027 | Post-Seed EBITDA milestone; ISO 27001 certified"),
    ("Series A | $3,000,000 | Month 30 (Q1 2029)", "Series A | $3,000,000 | Month 31 (December 2028)"),
    ("Break-even point | Month 28 (~323 paying clients)",
     "Cash-flow break-even | Month 28 (~323 clients) · EBITDA ~Month 18"),
    ("Secrets | AWS Secrets Manager | All API keys and credentials — zero hardcoded secrets",
     "Secrets | AWS SSM Parameter Store (MVP) → Secrets Manager (paid tier) | All API keys — zero hardcoded secrets"),
    ("CDN | AWS CloudFront | Frontend delivery + edge caching",
     "CDN | AWS Amplify (MVP) · CloudFront (paid tier) | Frontend delivery + edge caching"),
    # v1.15 — full blueprint alignment
    ("Lambda / Fargate / ECS infrastructure · Marketplace listing · ISV Accelerate co-sell",
     "API Gateway + Lambda + Amplify + SQS (serverless MVP) · Marketplace listing · ISV Accelerate co-sell"),
    ("6 months Lambda + EKS (paid) scan runtime + Lambda + S3 + Cognito",
     "12+ months serverless MVP (API GW + Lambda + Amplify + SQS) + S3 + Cognito"),
    ("8 / Platform", "8+L9 / Platform"),
    ("Stage:  Pre-Revenue / MVP", "Stage:  Pre-Revenue / Blueprint complete (socvault-app next)"),
    ("Founding Engineer (Python/DevSecOps) | Hiring — Immediate | FastAPI backend · scanner integration · AWS infrastructure · CI/CD pipeline",
     "Founding Engineer (Python/DevSecOps) | Hiring — Immediate | Lambda/API Gateway backend · scanner integration · Terraform · CI/CD to staging"),
    ("Month 15 (Q3 2027)", "Month 15 (August 2027)"),
    ("Super Admin API Explorer, encrypted credential vault, and live Development Tracker — built for MSPs and internal SecOps from day one.",
     "Super Admin API Explorer, Threat Intel registry (32 feeds), encrypted Pass & Keys vault, and Development Tracker — built for MSPs and internal SecOps from day one."),
    ("Target outcome: 50 paying clients · $10K MRR · Seed deck ready at Month 15.",
     "Target outcome: 50 paying clients · $10K MRR by ~M8 (Feb 2027); Seed deck ready; $750K Seed close at Month 15 (Aug 2027)."),
    ("Secrets in AWS SSM Parameter Store (MVP) → Secrets Manager at scale — zero credentials in source code",
     "Secrets in AWS SSM Parameter Store (MVP) → Secrets Manager (paid tier) — zero credentials in source code"),
    # v1.16 — deck re-audit (ISO timing, valuations, api-staging, 26 epics, chronology)
    ("208 user stories, 25 wireframes, OpenAPI",
     "208 user stories, 25 wireframes, 26 epics, OpenAPI"),
    ("208 user stories · 25 wireframes · OpenAPI",
     "208 user stories · 25 wireframes · 26 epics · OpenAPI"),
    ("208 user stories, 25 wireframes, OpenAPI, Threat Intel registry",
     "208 user stories, 25 wireframes, 26 epics, OpenAPI, Threat Intel registry"),
    ("208 user stories, 25 wireframes, OpenAPI, Threat Intel registry (32 feeds), serverless staging architecture (ADR-006)",
     "208 user stories, 25 wireframes, 26 epics, OpenAPI, Threat Intel registry (32 feeds), serverless staging architecture (ADR-006)"),
    ("Month 28 (~323 clients) · EBITDA ~Month 18",
     "Month 28 (September 2028, ~323 clients) · EBITDA ~Month 18"),
    ("cash-flow break-even Month 28, LTV:CAC = 12:1",
     "cash-flow break-even Month 28 (September 2028), LTV:CAC = 12:1"),
    ("Seed unlock criteria; SOAR live; Seed deck ready",
     "Seed unlock criteria; SOAR live; Seed deck ready; 3 MSP agreements signed"),
    ("AWS staging MVP · 50 paying clients · $10K MRR",
     "AWS staging MVP · 50 paying clients · $10K MRR · 3 MSP agreements · ISO in progress"),
]

STAGING_MVP_CRITERIA = (
    "api-staging.socvault.io green; L1 full 15-step scan completes in <3 min; "
    "Claude report generated; Bruno collection green"
)

CANONICAL_TIMELINE = (
    "Canonical financial timeline: Seed unlock 50/$10K ~M8 · Year 1 end 150 clients · "
    "Seed close ~200/~$16K M15 · M18 323/$22K EBITDA · cash-flow break-even M28 (Sep 2028)"
)

ASK_SUMMARY = (
    "Raising $251,000 on a SAFE note at a $2M valuation cap — 12-month runway to ship staging MVP on "
    "api-staging.socvault.io + app-staging.socvault.io, reach 50 paying clients and $10K MRR (~Month 8), "
    "and close a $750K Seed round ($3.5M pre-money) at Month 15 (August 2027). Product blueprint is "
    "complete (208 user stories, 25 wireframes, 26 epics, OpenAPI, ADR-006 serverless architecture); "
    "application repo (socvault-app) is the next build phase."
)

COVER_METRICS_STRIP = (
    "$251K Pre-Seed SAFE · $2M cap · Staging-first MVP (ADR-006) · 97.6% VAPT margin · "
    "8+L9 layers · 208 stories · 26 epics"
)

USP8_BODY = (
    "PCI-DSS 4.0, UK GDPR, ISO 27001:2022, SOC 2 Type II, and Cyber Essentials Plus included in the paid tier — "
    "not an add-on. No SMB security scanner includes all five frameworks. Compliance history accumulates as a "
    "switching cost that grows every month."
)

NEW_TRACTION_V15 = [
    "Runtime status: blueprint/documentation complete — no application code deployed yet; Phase 0 = AWS account + Terraform staging (see DEVELOPMENT_TRACKER.md)",
    CANONICAL_TIMELINE,
]

MILESTONE_SORT_PREFIXES = [
    "AWS staging MVP",
    "Production cutover live",
    "10 beta clients",
    "First paid conversion",
    "Product Hunt launch",
    "50 paying clients",
    "ISO 27001 Stage 1",
    "150 paying clients",
    "Seed round close",
    "ISO 27001 certified",
    "323 paying clients",
]

MILESTONE_ROW_FIXES = {
    "AWS + Azure Marketplace listings live": (
        "Seed round close",
        "August 2027 (Month 15)",
        "$750K raised · ~200 clients · ~$16K MRR · AWS Marketplace live",
    ),
    "ISO 27001 certification": (
        "ISO 27001 Stage 1 audit",
        "March 2027",
        "Zero critical findings",
    ),
    "MSP channel: 50+ clients via partners": (
        "150 paying clients (Year 1 end)",
        "June 2027",
        "$9K MRR · MSP channel foundation · ISO 27001 in progress",
    ),
}

MILESTONE_ROWS_TO_ADD = [
    ("Product Hunt launch", "December 2026", "Top 5 launch · organic inbound funnel"),
    ("ISO 27001 certified", "October 2027", "Certificate issued · enterprise sales enabled"),
]

PRESEED_UNLOCK = (
    "AWS staging MVP · 50 paying clients · $10K MRR · 3 MSP agreements · ISO in progress"
)
SEED_CLOSE_MILESTONE = (
    "~200 clients · ~$16K MRR at Seed close · AWS Marketplace live · "
    "ISO in progress (cert Oct 2027)"
)

REVENUE_SNAPSHOT_ROWS = [
    ("M6 (Nov 2026)", "20", "$1,100", "$13,200"),
    ("M12 (May 2027)", "150", "$9,000", "$108,000"),
    ("M15 Seed close (Aug 2027)", "~200", "~$16,000", "~$192,000"),
    ("M18 EBITDA (Nov 2027)", "323", "$22,000", "$264,000"),
]

TRACTION_UPDATES = {
    "Complete product blueprint: technical architecture, requirements, financial model, roadmap, brand — all documented":
        "Complete product blueprint: 208 user stories, 25 wireframes, OpenAPI, Threat Intel registry (32 feeds), serverless staging architecture (ADR-006)",
    "Complete product blueprint: 201 user stories, 25 wireframes, OpenAPI MVP contract, Bruno/Postman collections, financial model, CI/CD guide — all documented":
        "Complete product blueprint: 208 user stories, 25 wireframes, OpenAPI, Threat Intel registry (32 feeds), serverless staging architecture (ADR-006)",
    "Complete product blueprint: 208 user stories, 25 wireframes, OpenAPI contract, Threat Intel registry (32 feeds), CI/CD guide — all documented":
        "Complete product blueprint: 208 user stories, 25 wireframes, OpenAPI, Threat Intel registry (32 feeds), serverless staging architecture (ADR-006)",
    "8-layer scanning architecture designed and validated against industry tooling (Wazuh, Nuclei, Semgrep, MobSF, CloudFox)":
        "8 core scanning layers + L9 AI Agent Scan designed and validated against industry tooling (Wazuh, Nuclei, Semgrep, MobSF, CloudFox)",
    "Claude AI integration designed with correct model IDs (claude-sonnet-4-6), prompt caching architecture, and COGS model":
        "Claude AI integration designed with correct model IDs (claude-sonnet-4-6), prompt caching architecture, and fully-loaded COGS model (~$0.36/VAPT)",
}

NEW_TRACTION_BULLETS = [
    "AWS-only delivery: staging-first serverless MVP (API Gateway + Lambda + Amplify + SQS); production dormant until cutover (ADR-006)",
    "Infrastructure as Code (Terraform) + automated staging QA (`tests/qa/`) on every deploy; user stories built one-at-a-time",
    "Super Admin tooling: API Explorer, Pass & Keys vault, Threat Intel feed registry (32 free APIs), Development Tracker UI",
    "Paid-tier path documented: CloudFront CDN, AWS WAF, Backup, Amazon EKS (Kubernetes), AWS CodePipeline",
]

NEW_USP_BULLETS = [
    ("USP 9 — AI Chat Assistant", "Credit-based conversational AI that dispatches scan and remediation actions from natural language — 83% blended gross margin on credit bundles."),
    ("USP 10 — Operator Platform", "Super Admin API Explorer, Threat Intel feed registry (32 free APIs), encrypted credential vault, and Development Tracker — built for MSPs and internal SecOps."),
]

DEDUP_MARKERS = [
    "USP 9 — AI Chat Assistant",
    "USP 10 — Operator Platform",
    "Credit-based conversational AI that dispatches scan",
    "Super Admin API Explorer",
    "Super Admin tooling",
    "AWS-only delivery",
    "Infrastructure as Code",
    "Paid-tier path",
    "Threat Intel feed registry",
    "32 free APIs",
    "auto-scaling ECS Fargate",
    "Platform extensions",
]

# Paragraphs to remove entirely (stale duplicates after v1.14)
REMOVE_PARAGRAPH_PREFIXES = [
    "AWS-only delivery model: single Free Tier account",
    "Platform extensions specified:",
    "Super Admin tooling specified:",
]


def replace_in_text(text: str) -> str:
    for old, new in REPLACEMENTS:
        if old in text:
            text = text.replace(old, new)
    return text


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_text


def insert_paragraph_after(paragraph, text: str, style=None) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    replace_paragraph_text(new_para, text)
    return new_para


def full_doc_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def deduplicate_paragraphs(doc) -> int:
    seen = set()
    removed = 0
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        if any(t.startswith(prefix) for prefix in REMOVE_PARAGRAPH_PREFIXES):
            p._element.getparent().remove(p._element)
            removed += 1
            continue
        if any(t.startswith(m) or m in t for m in DEDUP_MARKERS):
            if t in seen:
                p._element.getparent().remove(p._element)
                removed += 1
            else:
                seen.add(t)
    return removed


def delete_table_row(row) -> None:
    row._tr.getparent().remove(row._tr)


def fill_ask_section(doc) -> None:
    """Section 1 'The Ask' — insert or sync summary paragraph."""
    body = full_doc_text(doc)
    ask_filled = False
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "The Ask":
            if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i + 1].text.strip():
                replace_paragraph_text(doc.paragraphs[i + 1], ASK_SUMMARY)
            elif ASK_SUMMARY[:40] not in body:
                insert_paragraph_after(p, ASK_SUMMARY)
            ask_filled = True
            break
    for p in doc.paragraphs:
        if p.text.strip().startswith("Raising $251,000"):
            if p.text.strip() != ASK_SUMMARY:
                replace_paragraph_text(p, ASK_SUMMARY)
            ask_filled = True
            break
    if not ask_filled:
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == "The Ask":
                insert_paragraph_after(p, ASK_SUMMARY)
                break
    if doc.tables:
        cover = doc.tables[0].rows[0].cells[0].text.strip()
        if not cover or "26 epic" not in cover.lower() or "208 stories" not in cover:
            replace_paragraph_text(
                doc.tables[0].rows[0].cells[0].paragraphs[0],
                COVER_METRICS_STRIP,
            )


def fix_usp8_compliance(doc) -> None:
    """Ensure USP 8 has compliance body; remove orphaned duplicate after USP 9/10."""
    usp8 = orphan = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "USP 8 — Compliance as Standard":
            usp8 = p
        if t.startswith("PCI-DSS 4.0, UK GDPR"):
            orphan = p
    if usp8 is None:
        return
    # Insert body if next paragraph is USP 9
    for p in doc.paragraphs:
        if p.text.strip().startswith("USP 9 — AI Chat"):
            if USP8_BODY not in full_doc_text(doc):
                insert_paragraph_after(usp8, USP8_BODY)
            break
    if orphan is not None and orphan._p.getparent() is not None:
        orphan._p.getparent().remove(orphan._p)


def add_traction_v15_bullets(doc) -> None:
    body = full_doc_text(doc)
    anchor = None
    for p in doc.paragraphs:
        if "Incorporated: SOCVault Ltd" in p.text:
            anchor = p
            break
    if anchor is None:
        for p in doc.paragraphs:
            if "Complete product blueprint: 208 user stories" in p.text:
                anchor = p
                break
    if anchor is None:
        return
    for bullet in NEW_TRACTION_V15:
        if bullet not in body:
            anchor = insert_paragraph_after(anchor, bullet)
            body = full_doc_text(doc)


def fix_milestone_table(doc) -> None:
    """Milestone roadmap: cutover, timing, M18, remove duplicates, add investor §16 rows."""
    for table in doc.tables:
        if len(table.rows) < 2 or table.rows[0].cells[0].text.strip() != "Milestone":
            continue
        # Remove duplicate Seed-unlock row
        for row in list(table.rows[1:]):
            if row.cells[0].text.strip().startswith("Seed unlock (see row above)"):
                delete_table_row(row)

        existing = {row.cells[0].text.strip() for row in table.rows[1:]}
        for row in table.rows[1:]:
            if len(row.cells) < 3:
                continue
            milestone = row.cells[0].text.strip()
            if milestone.startswith("AWS staging MVP"):
                replace_paragraph_text(row.cells[2].paragraphs[0], STAGING_MVP_CRITERIA)
            elif milestone.startswith("Production cutover live"):
                replace_paragraph_text(
                    row.cells[2].paragraphs[0],
                    "api.socvault.io + app.socvault.io live; cutover checklist signed off; post-cutover smoke QA green",
                )
            elif milestone.startswith("50 paying clients"):
                replace_paragraph_text(row.cells[0].paragraphs[0], "50 paying clients · $10K MRR")
                replace_paragraph_text(row.cells[1].paragraphs[0], "February 2027 (~M8)")
                replace_paragraph_text(
                    row.cells[2].paragraphs[0],
                    "Seed unlock criteria; SOAR live; Seed deck ready",
                )
            elif "$50K MRR / Series A ready" in milestone or milestone.startswith("$50K MRR"):
                replace_paragraph_text(row.cells[0].paragraphs[0], "323 paying clients · $22K MRR (M18)")
                replace_paragraph_text(row.cells[1].paragraphs[0], "November 2027 (M18)")
                replace_paragraph_text(
                    row.cells[2].paragraphs[0],
                    "Post-Seed EBITDA milestone · ISO 27001 certified",
                )
            elif milestone in MILESTONE_ROW_FIXES:
                m, d, s = MILESTONE_ROW_FIXES[milestone]
                replace_paragraph_text(row.cells[0].paragraphs[0], m)
                replace_paragraph_text(row.cells[1].paragraphs[0], d)
                replace_paragraph_text(row.cells[2].paragraphs[0], s)
                existing.discard(milestone)
                existing.add(m)

        for m, d, s in MILESTONE_ROWS_TO_ADD:
            if m not in existing:
                new_row = table.add_row()
                replace_paragraph_text(new_row.cells[0].paragraphs[0], m)
                replace_paragraph_text(new_row.cells[1].paragraphs[0], d)
                replace_paragraph_text(new_row.cells[2].paragraphs[0], s)

        sort_milestone_table_rows(table)
        break


def milestone_sort_key(milestone: str) -> int:
    for i, prefix in enumerate(MILESTONE_SORT_PREFIXES):
        if milestone.startswith(prefix):
            return i
    return len(MILESTONE_SORT_PREFIXES)


def sort_milestone_table_rows(table) -> None:
    """Reorder milestone rows chronologically (investor §16 sequence)."""
    rows_data = []
    for row in list(table.rows[1:]):
        rows_data.append(
            (
                milestone_sort_key(row.cells[0].text.strip()),
                row.cells[0].text.strip(),
                row.cells[1].text.strip(),
                row.cells[2].text.strip(),
            )
        )
        delete_table_row(row)
    rows_data.sort(key=lambda x: x[0])
    for _, m, d, s in rows_data:
        new_row = table.add_row()
        replace_paragraph_text(new_row.cells[0].paragraphs[0], m)
        replace_paragraph_text(new_row.cells[1].paragraphs[0], d)
        replace_paragraph_text(new_row.cells[2].paragraphs[0], s)


def add_revenue_snapshot_rows(doc) -> None:
    """Add M6/M12/M15/M18 rows to Year | Clients revenue table."""
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        h = [c.text.strip() for c in table.rows[0].cells]
        if h[0] != "Year" or "MRR" not in h:
            continue
        existing = {row.cells[0].text.strip() for row in table.rows[1:]}
        for label, clients, mrr, arr in REVENUE_SNAPSHOT_ROWS:
            if label in existing:
                continue
            row = table.add_row()
            vals = [label, clients, mrr, arr, "—"]
            for cell, val in zip(row.cells, vals[: len(row.cells)]):
                replace_paragraph_text(cell.paragraphs[0], val)
        break


def enhance_unit_economics_table(doc) -> None:
    for table in doc.tables:
        if len(table.rows) < 2 or table.rows[0].cells[0].text.strip() != "Metric":
            continue
        existing = {row.cells[0].text.strip() for row in table.rows[1:]}
        extras = [
            ("MRR at M18 (323 clients)", "$22,000"),
            ("Seed close M15 (~200 clients)", "~$16,000 MRR"),
        ]
        for metric, value in extras:
            if metric in existing:
                continue
            row = table.add_row()
            replace_paragraph_text(row.cells[0].paragraphs[0], metric)
            replace_paragraph_text(row.cells[1].paragraphs[0], value)
        break


def ensure_architecture_mvp_rows(doc) -> None:
    """Ensure API Gateway, SQS, DynamoDB appear in AWS architecture table."""
    want = [
        ("API (MVP)", "API Gateway HTTP + Lambda", "Cognito JWT authorizer · FastAPI/Mangum handlers · staging primary"),
        ("Async queue", "Amazon SQS + Lambda", "Scan jobs; replaces Celery on MVP (ADR-006)"),
        ("Telemetry / rate limits", "Amazon DynamoDB", "Per-tenant COGS, rate limits, vault sessions (TTL)"),
        ("Primary database", "MongoDB Atlas M0", "Document-per-tenant; Motor async driver on AWS eu-west-2"),
    ]
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        if table.rows[0].cells[0].text.strip() != "Component":
            continue
        if "AWS Service" not in table.rows[0].cells[1].text:
            continue
        existing = {row.cells[0].text.strip().lower() for row in table.rows[1:]}
        insert_after = table.rows[-1]
        for comp, svc, purpose in want:
            key = comp.split("(")[0].strip().lower()
            if any(key in e for e in existing):
                continue
            new_row = table.add_row()
            replace_paragraph_text(new_row.cells[0].paragraphs[0], comp)
            replace_paragraph_text(new_row.cells[1].paragraphs[0], svc)
            replace_paragraph_text(new_row.cells[2].paragraphs[0], purpose)
        break


def fix_financial_summary_table(doc) -> None:
    """Unit economics table: EBITDA vs cash break-even labels."""
    for table in doc.tables:
        if len(table.rows) < 2 or table.rows[0].cells[0].text.strip() != "Metric":
            continue
        for row in table.rows[1:]:
            label = row.cells[0].text.strip()
            if label.startswith("Break-even") or label.startswith("Cash-flow break-even"):
                replace_paragraph_text(row.cells[0].paragraphs[0], "Cash-flow break-even")
                replace_paragraph_text(
                    row.cells[1].paragraphs[0],
                    "Month 28 (September 2028, ~323 clients) · EBITDA ~Month 18",
                )
        break


def sync_traction_epic_count(doc) -> None:
    """Ensure traction bullets cite 26 epics alongside 208 stories."""
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Canonical financial timeline:"):
            replace_paragraph_text(p, CANONICAL_TIMELINE)
            continue
        if "208 user stories" in t and "26 epic" not in t.lower():
            updated = t.replace(
                "208 user stories, 25 wireframes, OpenAPI, Threat Intel",
                "208 user stories, 25 wireframes, 26 epics, OpenAPI, Threat Intel",
            )
            if updated == t:
                updated = t.replace(
                    "208 user stories, 25 wireframes, OpenAPI",
                    "208 user stories, 25 wireframes, 26 epics, OpenAPI",
                )
            if updated != t:
                replace_paragraph_text(p, updated)


def fix_funding_rounds_table(doc) -> None:
    for table in doc.tables:
        if len(table.rows) < 2 or table.rows[0].cells[0].text.strip() != "Round":
            continue
        for row in table.rows[1:]:
            name = row.cells[0].text.strip()
            if name.startswith("Pre-Seed") and len(row.cells) >= 4:
                replace_paragraph_text(row.cells[3].paragraphs[0], PRESEED_UNLOCK)
            elif name.startswith("Seed"):
                if "$3.5M" not in row.cells[1].text:
                    replace_paragraph_text(row.cells[1].paragraphs[0], "$750,000 ($3.5M pre-money)")
                if len(row.cells) >= 4:
                    replace_paragraph_text(row.cells[3].paragraphs[0], SEED_CLOSE_MILESTONE)
            elif name.startswith("Series A"):
                if "$12M" not in row.cells[1].text:
                    replace_paragraph_text(row.cells[1].paragraphs[0], "$3,000,000 ($12M pre-money)")
                replace_paragraph_text(row.cells[2].paragraphs[0], "Month 31 (December 2028)")
                if len(row.cells) >= 4:
                    replace_paragraph_text(
                        row.cells[3].paragraphs[0],
                        "2,000 clients · $150K MRR · US expansion · Managed SOC service",
                    )
        break


def fix_architecture_secrets_row(doc) -> None:
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        if table.rows[0].cells[0].text.strip() != "Component":
            continue
        if "AWS Service" not in table.rows[0].cells[1].text:
            continue
        for row in table.rows[1:]:
            comp = row.cells[0].text.strip()
            if comp == "Secrets" and "Secrets Manager" in row.cells[1].text and "SSM" not in row.cells[1].text:
                replace_paragraph_text(
                    row.cells[1].paragraphs[0],
                    "AWS SSM Parameter Store (MVP) → Secrets Manager (paid tier)",
                )
            if comp == "CDN" and "CloudFront" in row.cells[1].text and "Amplify" not in row.cells[1].text:
                replace_paragraph_text(
                    row.cells[1].paragraphs[0],
                    "AWS Amplify (MVP) · CloudFront (paid tier)",
                )
        break


def fix_table_cells(doc) -> None:
    for table in doc.tables:
        for row in table.rows:
            cells_text = [c.text for c in row.cells]
            joined = " | ".join(cells_text)
            if "COGS per Web VAPT" in joined or ("Web VAPT" in joined and "$0.052" in joined):
                for c in row.cells:
                    if "$0.052" in c.text and "malware" not in joined.lower():
                        c.text = c.text.replace("$0.052", "$0.36")
            if "L2" in joined and "$0.052" in joined and "VAPT" in joined:
                for c in row.cells:
                    c.text = c.text.replace("$0.052", "$0.36")
            if "Mobile" in joined and "97.6%" in joined:
                for c in row.cells:
                    c.text = c.text.replace("97.6%", "97.8%")
            if "Cloud" in joined and "97.6%" in joined and "CloudFox" in joined:
                for c in row.cells:
                    c.text = c.text.replace("97.6%", "97.8%")
            if "Celery" in joined and "Redis" in joined:
                for c in row.cells:
                    c.text = c.text.replace("Celery + Redis", "SQS + Lambda")
                    c.text = c.text.replace("Celery / Redis", "SQS + Lambda")
                    c.text = c.text.replace("Celery workers", "SQS + Lambda workers")
            if "ECS Fargate" in joined or "EC2 Docker" in joined or "Vercel" in joined:
                for c in row.cells:
                    c.text = c.text.replace("ECS Fargate", "API Gateway + Lambda")
                    c.text = c.text.replace("EC2 Docker", "serverless staging")
                    c.text = c.text.replace("Vercel", "Amplify")
            if "Fargate ECS" in joined or "Fargate (ECS)" in joined:
                for c in row.cells:
                    c.text = c.text.replace("Fargate ECS", "Lambda + EKS (paid)")
                    c.text = c.text.replace("AWS Fargate (ECS)", "Lambda / Fargate / EKS")
                    c.text = c.text.replace("Fargate scan tasks", "Serverless MVP + scan workers")
                    c.text = c.text.replace("6 months Fargate ECS scan runtime", "12+ months serverless MVP + Lambda")
            if "Scanning Runtime" in joined and "Fargate" in joined:
                for c in row.cells:
                    if "AWS Lambda (L1" in c.text:
                        c.text = "AWS Lambda (L1) · API Gateway + Lambda (MVP API) · EKS (paid scale) · EC2 (L7 Wazuh)"
            if "L2–L6 Paid Scans" in joined or "L2-L6 Paid Scans" in joined:
                for c in row.cells:
                    if "Fargate" in c.text:
                        c.text = "Lambda / Fargate / EKS (isolated per scan)"
            if "AWS Infrastructure (12 months)" in joined and "Fargate" in joined:
                for c in row.cells:
                    c.text = c.text.replace("Fargate scan tasks · Lambda", "API Gateway · Lambda · Amplify · SQS")
            if "AWS staging MVP" in joined and "Bruno" not in joined:
                for c in row.cells:
                    if "Claude report generated" in c.text and "Bruno" not in c.text:
                        c.text = c.text.replace(
                            "Claude report generated",
                            "Claude report generated; automated QA green",
                        )
            if "Cloudflare" in joined and "WAF" in joined:
                for c in row.cells:
                    c.text = c.text.replace("Cloudflare WAF · DDoS · CDN", "CloudFront · AWS WAF · CDN (paid tier)")
                    c.text = c.text.replace("Cloudflare WAF", "CloudFront + AWS WAF")
            if "Multi-tenant user pools" in joined or "user pools with custom tenant_id" in joined:
                for c in row.cells:
                    if "user pool" in c.text.lower():
                        c.text = "One Cognito User Pool per environment; tenant_id custom attribute + SMS OTP via SNS"
            if "dual-env single account" in joined:
                for c in row.cells:
                    c.text = c.text.replace("dual-env single account", "single account · staging active · production dormant")
            if "Production cutover live" in joined and "production dormant until cutover checklist" in joined:
                for c in row.cells:
                    if "production dormant until cutover checklist" in c.text:
                        c.text = "api.socvault.io + app.socvault.io live; cutover checklist signed off; post-cutover smoke QA green"


def ensure_l9_row(doc) -> None:
    for table in doc.tables:
        if len(table.rows) >= 2 and table.rows[0].cells[0].text.strip() == "Layer":
            if any(row.cells[0].text.strip() == "L9" for row in table.rows):
                return
            new_row = table.add_row()
            vals = [
                "L9", "AI Agent Scan", "Claude Opus (extended thinking)",
                "Autonomous multi-step security assessment with live activity log; OWASP-mapped findings; 1 scan/7 days",
                "SOC Pro (Phase 2)",
            ]
            for cell, val in zip(new_row.cells, vals):
                replace_paragraph_text(cell.paragraphs[0], val)


def process_doc():
    doc = Document(DECK)
    removed = deduplicate_paragraphs(doc)

    for p in doc.paragraphs:
        if p.text.strip():
            updated = replace_in_text(p.text)
            updated = TRACTION_UPDATES.get(updated, updated)
            if updated != p.text:
                replace_paragraph_text(p, updated)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        updated = replace_in_text(p.text)
                        if updated != p.text:
                            replace_paragraph_text(p, updated)

    fix_table_cells(doc)
    fill_ask_section(doc)
    fix_usp8_compliance(doc)
    fix_milestone_table(doc)
    fix_financial_summary_table(doc)
    enhance_unit_economics_table(doc)
    add_revenue_snapshot_rows(doc)
    fix_funding_rounds_table(doc)
    fix_architecture_secrets_row(doc)
    ensure_architecture_mvp_rows(doc)
    add_traction_v15_bullets(doc)
    sync_traction_epic_count(doc)
    ensure_l9_row(doc)

    body = full_doc_text(doc)

    if "Freemium 15-step L1 Recon process defined" in body:
        for p in doc.paragraphs:
            if "Freemium 15-step L1 Recon process defined" in p.text:
                anchor = p
                for bullet in reversed(NEW_TRACTION_BULLETS):
                    if bullet not in body:
                        anchor = insert_paragraph_after(anchor, bullet)
                break

    body = full_doc_text(doc)
    if "USP 8 — Compliance as Standard" in body and "USP 9 — AI Chat Assistant" not in body:
        for p in doc.paragraphs:
            if p.text.startswith("USP 8 — Compliance as Standard"):
                anchor = p
                for title, desc in NEW_USP_BULLETS:
                    anchor = insert_paragraph_after(anchor, title)
                    anchor = insert_paragraph_after(anchor, desc)
                break

    doc.save(DECK)
    print(f"Updated {DECK} (removed {removed} duplicate paragraph(s))")


if __name__ == "__main__":
    process_doc()
