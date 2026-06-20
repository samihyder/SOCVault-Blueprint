"""
SOCVault Revenue & OpEx Financial Model — DOCX Generator
Branding: SOCVault green (#39B54A) + dark teal (#1B4332) from actual logo assets
Output:   BusinessProposal/SOCVault-Revenue-OpEx-Model.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Brand palette (extracted from actual SOCVault logo) ───────────────────────
GREEN       = RGBColor(0x39, 0xB5, 0x4A)   # "SOC" bright green
DARK_TEAL   = RGBColor(0x1B, 0x43, 0x32)   # "Vault" dark teal / primary dark
GREEN_MID   = RGBColor(0x2E, 0x7D, 0x3E)   # mid-green for sub-headers
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
OFF_WHITE   = RGBColor(0xF9, 0xFB, 0xF9)
LIGHT_GREEN = RGBColor(0xE8, 0xF5, 0xE9)   # very light green row tint
MID_GREY    = RGBColor(0x6B, 0x74, 0x80)
DARK_GREY   = RGBColor(0x2D, 0x37, 0x48)
POSITIVE    = RGBColor(0xD1, 0xFA, 0xE5)   # green cell bg
NEGATIVE    = RGBColor(0xFE, 0xE2, 0xE2)   # red cell bg
AMBER_BG    = RGBColor(0xFE, 0xF3, 0xC7)
BLUE_BG     = RGBColor(0xDB, 0xEA, 0xFE)

LOGO_PATH        = "LogoAssets/SOCVault-Logo-Transparent.png"
LOGO_ICON_PATH   = "LogoAssets/SOCVault-Icon-128.png"


# ── XML helpers ───────────────────────────────────────────────────────────────

def hex_from_rgb(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, colour):
    """colour: RGBColor or hex string like '1B4332'"""
    h = colour if isinstance(colour, str) else hex_from_rgb(colour)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  h)
    tcPr.append(shd)


def set_cell_borders(cell, colour='C3E6CB', size='4'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    size)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), colour)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_padding(cell, top=60, bottom=60, left=80, right=80):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(existing)
    mar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def cell_para(cell, text, bold=False, size=9, colour=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False, space_b=30, space_a=30):
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if colour:
        run.font.color.rgb = colour
    return run


def add_para(doc, text='', bold=False, size=10, colour=None,
             align=WD_ALIGN_PARAGRAPH.LEFT, italic=False,
             space_before=3, space_after=5):
    p   = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = 'Calibri'
        if colour:
            run.font.color.rgb = colour
    return p


def add_divider(doc, colour='39B54A'):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), colour)
    pBdr.append(bot)
    pPr.append(pBdr)


def section_banner(doc, title, subtitle=''):
    """Dark teal full-width section header with green left accent."""
    tbl = doc.add_table(rows=1 if not subtitle else 2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Left accent strip
    tbl.columns[0].width = Inches(0.12)
    tbl.columns[1].width = Inches(6.88)

    def fill_banner_row(row, main_text, main_size, text_col, bg_col, accent_col):
        set_cell_bg(row.cells[0], accent_col)
        set_cell_bg(row.cells[1], bg_col)
        set_cell_borders(row.cells[0], colour=accent_col)
        set_cell_borders(row.cells[1], colour=bg_col)
        set_cell_padding(row.cells[0], 100, 100, 0, 0)
        set_cell_padding(row.cells[1], 80, 80, 120, 80)
        cell_para(row.cells[0], '')
        cell_para(row.cells[1], main_text, bold=True, size=main_size, colour=text_col)

    fill_banner_row(tbl.rows[0], title, 13, WHITE, '1B4332', '39B54A')
    if subtitle:
        fill_banner_row(tbl.rows[1], subtitle, 8.5, LIGHT_GREEN, '1B4332', '39B54A')

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def make_table(doc, headers, rows_data, col_widths,
               header_bg='1B4332', header_text=WHITE,
               alt_bg=LIGHT_GREEN, border_col='C3E6CB',
               font_size=8.5, bold_col0=True):
    """Build a styled table. rows_data: list of (values_tuple, optional_override_bg)."""
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(col_widths):
        for row in tbl.rows:
            row.cells[i].width = w

    # Header
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, header_bg)
        set_cell_borders(c, colour=header_bg)
        set_cell_padding(c, 80, 80, 80, 80)
        cell_para(c, h, bold=True, size=font_size, colour=header_text,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, row_item in enumerate(rows_data):
        if isinstance(row_item, tuple) and len(row_item) == 2 and isinstance(row_item[1], str):
            values, override_bg = row_item
        else:
            values, override_bg = row_item, None

        shade = ri % 2 == 0
        default_bg = hex_from_rgb(alt_bg) if shade else 'FFFFFF'
        bg = override_bg if override_bg else default_bg

        for ci, val in enumerate(values):
            c = tbl.rows[ri + 1].cells[ci]
            set_cell_bg(c, bg)
            set_cell_borders(c, colour=border_col)
            set_cell_padding(c, 60, 60, 80, 80)
            al = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            cell_para(c, str(val), bold=(bold_col0 and ci == 0),
                      size=font_size, align=al)

    return tbl


def total_row(tbl, row_idx, values, bg='1B4332', text_col=WHITE, font_size=8.5):
    row = tbl.rows[row_idx]
    for i, val in enumerate(values):
        c = row.cells[i]
        set_cell_bg(c, bg)
        set_cell_borders(c, colour=bg)
        set_cell_padding(c, 80, 80, 80, 80)
        al = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        cell_para(c, str(val), bold=True, size=font_size, colour=text_col, align=al)


# ── Document header & footer ──────────────────────────────────────────────────

def setup_header_footer(doc):
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        htbl = header.add_table(rows=1, cols=2, width=Inches(7.0))
        htbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        htbl.columns[0].width = Inches(2.5)
        htbl.columns[1].width = Inches(4.5)

        # Logo cell
        lc = htbl.rows[0].cells[0]
        set_cell_bg(lc, 'FFFFFF')
        set_cell_borders(lc, colour='FFFFFF')
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after  = Pt(0)
        if os.path.exists(LOGO_PATH):
            lr = lp.add_run()
            lr.add_picture(LOGO_PATH, width=Inches(1.5))

        # Title cell
        tc2 = htbl.rows[0].cells[1]
        set_cell_bg(tc2, 'FFFFFF')
        set_cell_borders(tc2, colour='FFFFFF')
        tp = tc2.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after  = Pt(0)
        tr1 = tp.add_run('Revenue & Operating Cost Model')
        tr1.bold = True
        tr1.font.size = Pt(8)
        tr1.font.color.rgb = DARK_TEAL
        tr1.font.name = 'Calibri'
        tp.add_run('\n')
        tr2 = tp.add_run('CONFIDENTIAL  ·  June 2026')
        tr2.font.size = Pt(7)
        tr2.font.color.rgb = MID_GREY
        tr2.font.name = 'Calibri'

        # Header bottom border
        hbdr_p = header.add_paragraph()
        hbdr_p.paragraph_format.space_before = Pt(2)
        hbdr_p.paragraph_format.space_after  = Pt(0)
        pPr  = hbdr_p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '39B54A')
        pBdr.append(bot)
        pPr.append(pBdr)

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        ftbl = footer.add_table(rows=1, cols=2, width=Inches(7.0))
        ftbl.columns[0].width = Inches(4.5)
        ftbl.columns[1].width = Inches(2.5)

        fc1 = ftbl.rows[0].cells[0]
        set_cell_bg(fc1, 'FFFFFF')
        set_cell_borders(fc1, colour='FFFFFF')
        fp1 = fc1.paragraphs[0]
        fp1.paragraph_format.space_before = Pt(4)
        r_f1 = fp1.add_run('SOCVault Ltd  ·  Registered in England & Wales  ·  socvault.io')
        r_f1.font.size = Pt(7)
        r_f1.font.color.rgb = MID_GREY
        r_f1.font.name = 'Calibri'

        fc2 = ftbl.rows[0].cells[1]
        set_cell_bg(fc2, 'FFFFFF')
        set_cell_borders(fc2, colour='FFFFFF')
        fp2 = fc2.paragraphs[0]
        fp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp2.paragraph_format.space_before = Pt(4)
        # Page number field
        r_pg = fp2.add_run('Page ')
        r_pg.font.size = Pt(7)
        r_pg.font.color.rgb = MID_GREY
        r_pg.font.name = 'Calibri'
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_el = fp2.add_run()
        run_el.font.size = Pt(7)
        run_el.font.color.rgb = MID_GREY
        run_el._r.append(fldChar1)
        run_el._r.append(instrText)
        run_el._r.append(fldChar2)


# ── Cover page ────────────────────────────────────────────────────────────────

def build_cover(doc):
    # Top green bar
    tbl_bar = doc.add_table(rows=1, cols=1)
    tbl_bar.columns[0].width = Inches(7.0)
    set_cell_bg(tbl_bar.rows[0].cells[0], '1B4332')
    set_cell_padding(tbl_bar.rows[0].cells[0], 160, 160, 80, 80)
    cell_para(tbl_bar.rows[0].cells[0], '')

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # Logo centred
    if os.path.exists(LOGO_PATH):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(0)
        p_logo.paragraph_format.space_after  = Pt(12)
        r_logo = p_logo.add_run()
        r_logo.add_picture(LOGO_PATH, width=Inches(3.2))

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after  = Pt(4)
    r_t = p_title.add_run('Revenue & Operating Cost\nFinancial Model')
    r_t.bold = True
    r_t.font.size = Pt(24)
    r_t.font.color.rgb = DARK_TEAL
    r_t.font.name = 'Calibri'

    # Green accent line under title
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_before = Pt(2)
    p_line.paragraph_format.space_after  = Pt(18)
    r_line = p_line.add_run('━━━━━━━━━━━━━━━━━━━━━━━━')
    r_line.font.color.rgb = GREEN
    r_line.font.size = Pt(10)

    # Metadata table
    meta = [
        ('Document Type',  'Internal Financial Model — Investor & Partner Reference'),
        ('Company',        'SOCVault Ltd — Registered in England & Wales'),
        ('Pakistan Entity','SOCVault (SMC-PVT) Ltd — Reg. BG96012, Karachi'),
        ('Stage',          'Pre-Revenue · Blueprint Complete · Pre-Seed Raise'),
        ('Model Period',   'Year 1 (2026–27) through Year 5 (2030–31)'),
        ('Team Structure', 'UK Anchor (Founders + CTO + BD) · Pakistan Engineering & Digital'),
        ('Currency',       'GBP primary · USD for Pakistan costs · rate ~0.79'),
        ('Prepared',       'June 2026  ·  SOCVault Founders'),
        ('Classification', 'CONFIDENTIAL — Not for public distribution'),
    ]

    tbl_meta = doc.add_table(rows=len(meta), cols=2)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_meta.columns[0].width = Inches(2.2)
    tbl_meta.columns[1].width = Inches(4.5)
    for i, (k, v) in enumerate(meta):
        shade = i % 2 == 0
        lc = tbl_meta.rows[i].cells[0]
        rc = tbl_meta.rows[i].cells[1]
        set_cell_bg(lc, '1B4332')
        set_cell_bg(rc, hex_from_rgb(LIGHT_GREEN) if shade else 'FFFFFF')
        set_cell_borders(lc, colour='39B54A')
        set_cell_borders(rc, colour='C3E6CB')
        set_cell_padding(lc, 70, 70, 90, 90)
        set_cell_padding(rc, 70, 70, 90, 90)
        cell_para(lc, k, bold=True, size=8.5, colour=GREEN)
        cell_para(rc, v, size=8.5, colour=DARK_GREY if i != len(meta)-1 else RGBColor(0xB4, 0x51, 0x09))

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Tagline
    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tag = p_tag.add_run(
        'The Unified AI-Enabled Cybersecurity Solution for SMBs\n'
        'Replacing a £28,000/year tool stack for £199/month'
    )
    r_tag.italic = True
    r_tag.font.size = Pt(9.5)
    r_tag.font.color.rgb = MID_GREY
    r_tag.font.name = 'Calibri'

    # Bottom teal bar
    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    tbl_bot = doc.add_table(rows=1, cols=3)
    tbl_bot.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (w, bg, txt) in enumerate([
        (Inches(2.3), '1B4332', '97.6% Gross Margin\nper VAPT Scan'),
        (Inches(2.3), '39B54A', '$9M ARR Target\nYear 5'),
        (Inches(2.3), '2E7D3E', '$251K Pre-Seed SAFE\n$2M Valuation Cap'),
    ]):
        c = tbl_bot.rows[0].cells[i]
        c.width = w
        set_cell_bg(c, bg)
        set_cell_borders(c, colour=bg)
        set_cell_padding(c, 120, 120, 80, 80)
        cell_para(c, txt, bold=True, size=9, colour=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()


# ── Section 1 – Executive Summary ────────────────────────────────────────────

def build_exec_summary(doc):
    section_banner(doc, '01   EXECUTIVE SUMMARY',
                   'Financial model rationale, team structure and headline metrics')

    add_para(doc,
        'SOCVault is a pre-revenue, AI-powered cybersecurity SaaS targeting UK SMBs (50–500 employees). '
        'The business operates a lean hybrid team model: a UK anchor team covering leadership, architecture, '
        'and business development; and a Pakistan-based engineering and digital team delivering the product '
        'at a fraction of UK market cost. The Pakistan entity (SOCVault SMC-PVT Ltd, Reg. BG96012, Karachi) '
        'employs the engineering team under local labour law.',
        size=9.5, colour=DARK_GREY, space_before=2, space_after=5)

    add_para(doc,
        'This document models revenue, gross margin, and operating expenditure across five reporting periods '
        '— from the first milestone at Month 6 through Year 5 — under three OpEx scenarios: Lean, Base, and '
        'Full Team. GBP figures assume a USD/GBP exchange rate of 0.79 (June 2026). Pakistan costs are shown '
        'in USD as that is the standard for offshore contractor rates in Pakistan\'s tech sector.',
        size=9.5, colour=DARK_GREY, space_before=2, space_after=8)

    add_para(doc, 'Key Financial Headlines', bold=True, size=11, colour=DARK_TEAL, space_before=4, space_after=3)

    bullets = [
        ('97.6%',  'Gross margin per VAPT scan — $0.36 COGS on $15.00 revenue (Claude + compute)'),
        ('92.5%',  'Gross margin SOC Pro — ~$15/month COGS on $199/month subscription'),
        ('82.7%',  'Year 1 blended gross margin (free-tier dilution in early months)'),
        ('95%+',   'Year 5 blended gross margin at scale with near-zero marginal COGS'),
        ('12:1',   'LTV:CAC ratio in Year 1 — $720 LTV vs $60 blended CAC'),
        ('M18',    'EBITDA break-even milestone (~250 clients, ~$22K MRR)'),
        ('M28',    'Cash-flow break-even (~323 clients)'),
        ('18–20mo','Effective runway on $251K raise under Lean Scenario (UK+Pakistan model)'),
    ]

    tbl = doc.add_table(rows=len(bullets), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Inches(0.9)
    tbl.columns[1].width = Inches(6.1)
    for i, (metric, label) in enumerate(bullets):
        shade = i % 2 == 0
        lc = tbl.rows[i].cells[0]
        rc = tbl.rows[i].cells[1]
        set_cell_bg(lc, '1B4332')
        set_cell_bg(rc, hex_from_rgb(LIGHT_GREEN) if shade else 'FFFFFF')
        set_cell_borders(lc, colour='39B54A')
        set_cell_borders(rc, colour='C3E6CB')
        set_cell_padding(lc, 60, 60, 80, 80)
        set_cell_padding(rc, 60, 60, 90, 90)
        cell_para(lc, metric, bold=True, size=9.5, colour=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_para(rc, label, size=9, colour=DARK_GREY)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_divider(doc)
    doc.add_page_break()


# ── Section 2 – Unit Economics ────────────────────────────────────────────────

def build_unit_economics(doc):
    section_banner(doc, '02   UNIT ECONOMICS',
                   'Per-transaction COGS, product-level margin and key ratios')

    add_para(doc, '2.1  Product Tier Gross Margins', bold=True, size=11,
             colour=DARK_TEAL, space_before=2, space_after=4)

    hdrs = ['Tier', 'Price / Month', 'COGS / Month', 'Gross Profit', 'Gross Margin']
    rows = [
        (('Free (Freemium)',         '$0',         '$0.03 / scan',  '$0',      '—'),        'E8F5E9'),
        (('Pay-Per-Target (VAPT)',   '$15 / IP',   '$0.36 / scan',  '$14.64',  '97.6%'),    'D1FAE5'),
        (('SOC Starter',             '$49 / month','~$5',           '~$44',    '~90.0%'),   'E8F5E9'),
        (('SOC Pro',                 '$199 / month','~$15',         '~$184',   '92.5%'),    'D1FAE5'),
        (('Enterprise',              '$499 / month','~$35',         '~$464',   '93.0%'),    'E8F5E9'),
        (('AI Chat Credits (blend)', 'Variable',   'Variable',      '—',       '~83.0%'),   'FFFFFF'),
    ]
    widths = [Inches(1.7), Inches(1.3), Inches(1.3), Inches(1.3), Inches(1.2)]
    tbl = make_table(doc, hdrs, rows, widths, font_size=8.5)
    # Colour margin column green
    for ri in range(1, len(rows) + 1):
        if tbl.rows[ri].cells[4].paragraphs[0].text not in ('—', ''):
            set_cell_bg(tbl.rows[ri].cells[4], 'D1FAE5')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_para(doc, '2.2  AI Cost Breakdown (Claude API with Prompt Caching)', bold=True, size=11,
             colour=DARK_TEAL, space_before=6, space_after=4)

    hdrs2 = ['Operation', 'Model', 'Input Tokens', 'Output Tokens', 'COGS / Call']
    rows2 = [
        ('L1 Recon AI report',    'claude-sonnet-4-6', '6,000 (cached)',  '800',   '$0.031'),
        ('L2–L6 VAPT scan',       'claude-sonnet-4-6', '12,000 (cached)', '1,500', '$0.360'),
        ('SOAR alert triage',     'claude-haiku-4-5',  '2,000 (cached)',  '300',   '$0.006'),
        ('L8 Malware analysis',   'claude-sonnet-4-6', '8,000 (cached)',  '1,200', '$0.042'),
        ('AI Chat (avg session)', 'claude-sonnet-4-6', '5,000 (cached)',  '1,000', '$0.027'),
    ]
    widths2 = [Inches(1.9), Inches(1.5), Inches(1.3), Inches(1.3), Inches(1.0)]
    make_table(doc, hdrs2, rows2, widths2, font_size=8.5)

    add_para(doc,
        'Pricing basis: claude-sonnet-4-6 = $3.00/1M input · $15.00/1M output · $0.30/1M cached reads. '
        'Prompt caching on all system prompts reduces effective input cost by ~90%. AI COGS is the smallest '
        'component of total cost at scale.',
        size=8, colour=MID_GREY, italic=True, space_before=4, space_after=8)

    add_para(doc, '2.3  Key Unit Economic Ratios (Year 1)', bold=True, size=11,
             colour=DARK_TEAL, space_before=4, space_after=4)

    metrics = [
        ('Blended CAC (Year 1)',             '$60'),
        ('Blended LTV (Year 1 — 12 months)', '$720'),
        ('LTV : CAC ratio',                  '12 : 1'),
        ('Payback period',                   '< 1 month on SOC Pro'),
        ('EBITDA break-even',                '~Month 18'),
        ('Cash-flow break-even',             'Month 28 (~323 clients)'),
        ('MRR at Month 18',                  '~$22,000'),
        ('MRR at Seed close (Month 15)',     '~$16,000'),
    ]
    tbl3 = doc.add_table(rows=len(metrics), cols=2)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.columns[0].width = Inches(3.5)
    tbl3.columns[1].width = Inches(3.0)
    for i, (k, v) in enumerate(metrics):
        shade = i % 2 == 0
        lc = tbl3.rows[i].cells[0]
        rc = tbl3.rows[i].cells[1]
        bg = hex_from_rgb(LIGHT_GREEN) if shade else 'FFFFFF'
        set_cell_bg(lc, bg); set_cell_bg(rc, bg)
        set_cell_borders(lc, 'C3E6CB'); set_cell_borders(rc, 'C3E6CB')
        set_cell_padding(lc, 70, 70, 90, 90); set_cell_padding(rc, 70, 70, 90, 90)
        cell_para(lc, k, bold=True, size=9, colour=DARK_TEAL)
        cell_para(rc, v, size=9, colour=DARK_GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_divider(doc)
    doc.add_page_break()


# ── Section 3 – Revenue Projections ──────────────────────────────────────────

def build_revenue(doc):
    section_banner(doc, '03   REVENUE PROJECTIONS',
                   'All milestones from Month 6 through Year 5')

    add_para(doc,
        'Revenue is modelled on a blended ARPU of ~$150/month at maturity, driven by the SOC Pro '
        '($199/month) tier as the primary retention product and Pay-Per-Target ($15/IP) for initial '
        'conversion. Free-tier users are included in total client counts but contribute $0 revenue, '
        'which depresses blended gross margin in Year 1 (82.7%) relative to mature state (95%+).',
        size=9.5, colour=DARK_GREY, space_before=2, space_after=6)

    add_para(doc, '3.1  Milestone Targets — Pre-Seed Phase', bold=True, size=11,
             colour=DARK_TEAL, space_before=2, space_after=4)

    hdrs = ['Milestone', 'Date', 'Paying Clients', 'MRR (USD)', 'ARR (USD)', 'Gross Margin']
    rows = [
        ('Staging MVP live',       'Month 3 (Sep 2026)',  '0',     '$0',       '$0',         '—'),
        ('First paying clients',   'Month 6 (Nov 2026)',  '20',    '$1,100',   '$13,200',    '~80%'),
        ('Pre-Seed milestone',     'Month 12 (May 2027)', '150',   '$9,000',   '$108,000',   '82.7%'),
        ('Seed close target',      'Month 15 (Aug 2027)', '~200',  '~$16,000', '~$192,000',  '~85%'),
        ('EBITDA break-even',      'Month 18 (Nov 2027)', '~250',  '~$22,000', '~$264,000',  '~89%'),
        ('Cash-flow break-even',   'Month 28 (Sep 2028)', '~323',  '~$30,000', '~$360,000',  '~91%'),
    ]
    widths = [Inches(1.7), Inches(1.5), Inches(1.1), Inches(1.0), Inches(1.0), Inches(1.0)]
    make_table(doc, hdrs, rows, widths, font_size=8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_para(doc, '3.2  Annual Revenue Summary — Year 1 to Year 5', bold=True, size=11,
             colour=DARK_TEAL, space_before=4, space_after=4)

    hdrs2 = ['Year', 'Period', 'Total Clients', 'MRR (USD)', 'ARR (USD)', 'ARR (GBP)', 'Gross Margin']
    rows2 = [
        ('Year 1', '2026–27', '150',    '$9,000',    '$108,000',   '£85,320',    '82.7%'),
        ('Year 2', '2027–28', '600',    '$45,000',   '$540,000',   '£426,600',   '91.0%'),
        ('Year 3', '2028–29', '2,000',  '$150,000',  '$1,800,000', '£1,422,000', '93.0%'),
        ('Year 4', '2029–30', '5,000',  '$375,000',  '$4,500,000', '£3,555,000', '94.0%'),
        ('Year 5', '2030–31', '10,000', '$750,000',  '$9,000,000', '£7,110,000', '95%+'),
    ]
    widths2 = [Inches(0.55), Inches(0.85), Inches(0.95), Inches(0.9), Inches(1.0), Inches(1.0), Inches(1.0)]
    tbl2 = make_table(doc, hdrs2, rows2, widths2, font_size=8.5)
    for ri in range(1, 6):
        set_cell_bg(tbl2.rows[ri].cells[6], 'D1FAE5')

    add_para(doc,
        'GBP conversion at USD/GBP 0.79. Year 4 is interpolated. Gross margin improves year-on-year '
        'as paying clients displace free-tier users and fixed infrastructure cost is amortised.',
        size=8, colour=MID_GREY, italic=True, space_before=4, space_after=8)

    add_para(doc, '3.3  Revenue Mix Shift by Year', bold=True, size=11,
             colour=DARK_TEAL, space_before=4, space_after=4)

    add_para(doc,
        'The mix shifts from Pay-Per-Target-heavy in Year 1 toward SOC Pro and Enterprise subscriptions '
        'in Year 3+. This is the primary driver of margin improvement — subscription COGS (~$15/month) '
        'are more predictable and lower than per-scan compute costs.',
        size=9.5, colour=DARK_GREY, space_before=2, space_after=5)

    hdrs3 = ['Year', 'Free Tier', 'Pay-Per-Target', 'SOC Pro', 'Enterprise', 'Blended ARPU']
    rows3 = [
        ('Year 1', '40%', '35%', '24%', '1%',  '~$60/mo'),
        ('Year 2', '25%', '25%', '45%', '5%',  '~$100/mo'),
        ('Year 3', '15%', '15%', '60%', '10%', '~$135/mo'),
        ('Year 4', '10%', '10%', '65%', '15%', '~$150/mo'),
        ('Year 5', '8%',  '7%',  '65%', '20%', '~$150/mo'),
    ]
    widths3 = [Inches(0.65), Inches(0.9), Inches(1.1), Inches(0.9), Inches(1.0), Inches(1.1)]
    make_table(doc, hdrs3, rows3, widths3, font_size=8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_divider(doc)
    doc.add_page_break()


# ── Section 4 – Team Structure & OpEx ────────────────────────────────────────

def build_opex(doc):
    section_banner(doc, '04   OPERATING EXPENDITURE MODEL',
                   'UK anchor + Pakistan engineering · detailed role-by-role cost across 5 years')

    add_para(doc,
        'SOCVault operates a deliberate geographic cost arbitrage: UK-based founders and one business '
        'development hire handle client relationships, investor relations, and product direction. All '
        'engineering, QA, DevOps, security tooling, and digital marketing is delivered from Pakistan '
        'through the Pakistan entity (SOCVault SMC-PVT Ltd). This reduces the annual salary burden '
        'by approximately 65–70% versus an equivalent all-UK team.',
        size=9.5, colour=DARK_GREY, space_before=2, space_after=8)

    # ── UK Team ──
    add_para(doc, '4.1  UK Team — Roles & Annual Cost (GBP, includes Employer NI at 13.8%)',
             bold=True, size=11, colour=DARK_TEAL, space_before=2, space_after=4)

    add_para(doc,
        'Pre-seed founders typically take below-market salaries or fully defer compensation to maximise '
        'runway. Employer National Insurance (13.8%) is included in all employed figures.',
        size=8, colour=MID_GREY, italic=True, space_before=0, space_after=5)

    hdrs_uk = ['Role', 'Year 1', 'Year 2', 'Year 3', 'Year 4', 'Year 5', 'Notes']
    rows_uk_data = [
        ('Founder / CEO',
         '£25,000', '£45,000', '£65,000', '£80,000', '£95,000',
         'Lean/deferred Y1; market rate Y3+'),
        ('Co-Founder / CTO / Solution Architect',
         '£28,000', '£50,000', '£70,000', '£85,000', '£100,000',
         'Technical co-founder; market rate Y3+'),
        ('Business Development Manager',
         '£28,000', '£32,000', '£38,000', '£45,000', '£52,000',
         'Commission element added from Y2'),
        ('UK Sales / Partnerships Manager',
         '—',       '—',       '£35,000', '£45,000', '£55,000',
         'New hire Y3 for EU market expansion'),
        ('Employer NI on UK salaries (~13.8%)',
         '£5,600',  '£9,300',  '£23,000', '£29,000', '£35,000',
         'Applies to all employed UK staff'),
    ]

    tbl_uk = doc.add_table(rows=1 + len(rows_uk_data) + 1, cols=7)
    tbl_uk.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_uk = [Inches(1.9), Inches(0.72), Inches(0.72), Inches(0.72), Inches(0.72), Inches(0.72), Inches(1.6)]
    for i, w in enumerate(widths_uk):
        for row in tbl_uk.rows:
            row.cells[i].width = w
    # Header
    for i, h in enumerate(hdrs_uk):
        c = tbl_uk.rows[0].cells[i]
        set_cell_bg(c, '1B4332'); set_cell_borders(c, '1B4332')
        set_cell_padding(c, 80, 80, 80, 80)
        cell_para(c, h, bold=True, size=8.5, colour=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    # Data
    for ri, r_ in enumerate(rows_uk_data):
        shade = ri % 2 == 0
        bg = hex_from_rgb(LIGHT_GREEN) if shade else 'FFFFFF'
        for ci, val in enumerate(r_):
            c = tbl_uk.rows[ri + 1].cells[ci]
            set_cell_bg(c, bg); set_cell_borders(c, 'C3E6CB')
            set_cell_padding(c, 65, 65, 80, 80)
            al = WD_ALIGN_PARAGRAPH.LEFT if ci in (0, 6) else WD_ALIGN_PARAGRAPH.CENTER
            cell_para(c, val, bold=(ci == 0), size=8.5, align=al)
    # Total row
    total_row(tbl_uk, len(rows_uk_data) + 1,
              ('UK Team Total (GBP)', '£86,600', '£136,300', '£231,000', '£284,000', '£337,000', ''))

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Pakistan Team ──
    add_para(doc, '4.2  Pakistan Engineering & Digital Team — Roles & Annual Cost (USD → GBP)',
             bold=True, size=11, colour=DARK_TEAL, space_before=6, space_after=4)

    add_para(doc,
        'Pakistan contractor rates are quoted in USD (local market standard for tech roles). '
        'Rates increase ~10–15%/year due to market competition. Converted to GBP at 0.79 for '
        'consolidated P&L. No employer NI equivalent for independent contractors.',
        size=8, colour=MID_GREY, italic=True, space_before=0, space_after=5)

    hdrs_pk = ['Role', 'Y1 (USD)', 'Y2 (USD)', 'Y3 (USD)', 'Y4 (USD)', 'Y5 (USD)', 'Notes']
    rows_pk_data = [
        ('Senior Backend Dev (Python / FastAPI)',
         '$13,200', '$14,800', '$18,000', '$21,000', '$24,000', 'Core API & scan engine'),
        ('Backend Developer #2',
         '$10,800', '$12,000', '$15,000', '$18,000', '$21,000', 'Scales to Senior by Y3'),
        ('Frontend Developer (React / Next.js)',
         '$9,600',  '$10,800', '$13,200', '$15,600', '$18,000', 'Dashboard & client portal'),
        ('DevOps / Cloud Engineer (AWS)',
         '$12,000', '$13,200', '$16,800', '$19,200', '$22,800', 'Infra, CI/CD, cost ops'),
        ('Security / Scanning Engineer',
         '$12,000', '$13,200', '$16,800', '$21,000', '$24,000', 'L1–L8 toolchain, Wazuh'),
        ('QA / Test Engineer',
         '$6,000',  '$7,200',  '$9,600',  '$12,000', '$14,400', 'Manual + automated testing'),
        ('Digital Marketing / SEO',
         '$6,000',  '$7,200',  '$8,400',  '$9,600',  '$10,800', 'Inbound, SEO, LinkedIn'),
        ('UI / UX Designer',
         '$4,800',  '$6,000',  '$8,400',  '$10,800', '$12,000', 'Design system, wireframes'),
        ('Backend Developer #3',
         '—',       '—',       '$13,200', '$15,600', '$18,000', 'Added Y3 for velocity'),
        ('Customer Success / Support',
         '—',       '$6,000',  '$7,200',  '$8,400',  '$9,600',  'Tier-1 support from Y2'),
    ]

    tbl_pk = doc.add_table(rows=1 + len(rows_pk_data) + 2, cols=7)
    tbl_pk.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_pk = [Inches(1.8), Inches(0.7), Inches(0.7), Inches(0.7), Inches(0.7), Inches(0.7), Inches(1.6)]
    for i, w in enumerate(widths_pk):
        for row in tbl_pk.rows:
            row.cells[i].width = w
    for i, h in enumerate(hdrs_pk):
        c = tbl_pk.rows[0].cells[i]
        set_cell_bg(c, '1B4332'); set_cell_borders(c, '1B4332')
        set_cell_padding(c, 80, 80, 80, 80)
        cell_para(c, h, bold=True, size=8.5, colour=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, r_ in enumerate(rows_pk_data):
        shade = ri % 2 == 0
        bg = hex_from_rgb(LIGHT_GREEN) if shade else 'FFFFFF'
        for ci, val in enumerate(r_):
            c = tbl_pk.rows[ri + 1].cells[ci]
            set_cell_bg(c, bg); set_cell_borders(c, 'C3E6CB')
            set_cell_padding(c, 65, 65, 80, 80)
            al = WD_ALIGN_PARAGRAPH.LEFT if ci in (0, 6) else WD_ALIGN_PARAGRAPH.CENTER
            cell_para(c, val, bold=(ci == 0), size=8.5, align=al)
    total_row(tbl_pk, len(rows_pk_data) + 1,
              ('Pakistan Total (USD)', '$74,400', '$90,400', '$126,600', '$151,200', '$174,600', ''))
    total_row(tbl_pk, len(rows_pk_data) + 2,
              ('Pakistan Total (GBP ~0.79)', '£58,776', '£71,416', '£100,014', '£119,448', '£137,934', ''),
              bg='2E7D3E')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Infrastructure & Tools ──
    add_para(doc, '4.3  Infrastructure, Tools & Overheads (GBP)',
             bold=True, size=11, colour=DARK_TEAL, space_before=6, space_after=4)

    hdrs_inf = ['Category', 'Year 1', 'Year 2', 'Year 3', 'Year 4', 'Year 5', 'Notes']
    rows_inf_data = [
        ('AWS Infrastructure (net of Activate credits)',
         '£3,000', '£8,000', '£22,000', '£55,000', '£130,000',
         'Lambda/SQS/EKS · credits offset Y1 heavily'),
        ('MongoDB Atlas',
         '£1,200', '£3,600', '£9,600', '£18,000', '£36,000',
         'M0 free → M10 → M30 at scale'),
        ('Third-party security APIs (Shodan, VirusTotal)',
         '£2,400', '£3,600', '£6,000', '£9,600', '£14,400',
         'Recon data feeds per scan volume'),
        ('SaaS tools (Notion, Slack, GitHub, Linear)',
         '£1,800', '£2,400', '£3,600', '£4,800', '£6,000',
         'Team collaboration & project mgmt'),
        ('Legal & compliance (UK)',
         '£3,500', '£5,000', '£8,000', '£12,000', '£18,000',
         'GDPR, contracts, company secretary'),
        ('Accountant / bookkeeping',
         '£2,000', '£2,500', '£3,500', '£5,000', '£7,000',
         'Annual accounts + management accounts'),
        ('Cyber & professional indemnity insurance',
         '£1,500', '£2,500', '£4,000', '£6,500', '£10,000',
         'E&O / PI — required for security product'),
        ('Marketing / PR / events',
         '£3,000', '£8,000', '£18,000', '£35,000', '£70,000',
         'Paid ads, events, analyst relations'),
        ('Equipment & software licences',
         '£2,500', '£1,500', '£2,000', '£3,000', '£4,000',
         'Laptops Y1; maintenance after'),
        ('Travel (UK client visits + Pakistan)',
         '£2,000', '£4,000', '£7,000', '£10,000', '£15,000',
         '1–2 PK visits/year + UK client meetings'),
        ('Contingency (5% of sub-total)',
         '£1,100', '£2,000', '£4,200', '£7,900', '£15,500',
         'Buffer for unplanned spend'),
    ]

    tbl_inf = doc.add_table(rows=1 + len(rows_inf_data) + 1, cols=7)
    tbl_inf.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_inf = [Inches(2.0), Inches(0.65), Inches(0.65), Inches(0.72), Inches(0.78), Inches(0.85), Inches(1.5)]
    for i, w in enumerate(widths_inf):
        for row in tbl_inf.rows:
            row.cells[i].width = w
    for i, h in enumerate(hdrs_inf):
        c = tbl_inf.rows[0].cells[i]
        set_cell_bg(c, '1B4332'); set_cell_borders(c, '1B4332')
        set_cell_padding(c, 80, 80, 80, 80)
        cell_para(c, h, bold=True, size=8.5, colour=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, r_ in enumerate(rows_inf_data):
        shade = ri % 2 == 0
        bg = hex_from_rgb(LIGHT_GREEN) if shade else 'FFFFFF'
        for ci, val in enumerate(r_):
            c = tbl_inf.rows[ri + 1].cells[ci]
            set_cell_bg(c, bg); set_cell_borders(c, 'C3E6CB')
            set_cell_padding(c, 65, 65, 80, 80)
            al = WD_ALIGN_PARAGRAPH.LEFT if ci in (0, 6) else WD_ALIGN_PARAGRAPH.CENTER
            cell_para(c, val, bold=(ci == 0), size=8.5, align=al)
    total_row(tbl_inf, len(rows_inf_data) + 1,
              ('Infrastructure & Overheads Total', '£24,000', '£43,100', '£87,900', '£166,800', '£325,900', ''))

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_divider(doc)
    doc.add_page_break()


# ── Section 5 – Three Scenarios ───────────────────────────────────────────────

def build_scenarios(doc):
    section_banner(doc, '05   OPEX SCENARIOS — LEAN / BASE / FULL TEAM',
                   'Three realistic cost structures: minimum viable to fully staffed')

    add_para(doc,
        'All three scenarios assume the same UK anchor structure (Founders + CTO + BD). '
        'The difference lies in Pakistan team size, hiring pace, and founder compensation level. '
        'Net P&L uses gross profit figures from Section 3.2 (converted at 0.79 USD/GBP). '
        'Negative figures are shown in parentheses.',
        size=9.5, colour=DARK_GREY, space_before=2, space_after=8)

    scenarios = [
        {
            'label':  'SCENARIO A — LEAN',
            'desc':   'Founders on minimum/deferred salary. Pakistan team limited to 4–5 core roles. '
                      'AWS Activate credits received ($25K–$100K). Best suited for Month 1–12 on $251K raise. '
                      'Extends effective runway to 18–20 months.',
            'bg':     '1B4332',
            'accent': '39B54A',
            'row_bg': 'E8F5E9',
            'years': [
                ('Year 1', '£95,000',  '£45,000',  '£18,000', '£158,000', '£70,000',  '(£88,000)'),
                ('Year 2', '£120,000', '£55,000',  '£32,000', '£207,000', '£337,140', '£130,140'),
                ('Year 3', '£160,000', '£78,000',  '£65,000', '£303,000', '£1,122,780','£819,780'),
                ('Year 4', '£210,000', '£95,000',  '£130,000','£435,000', '£2,807,700','£2,372,700'),
                ('Year 5', '£260,000', '£110,000', '£250,000','£620,000', '£5,616,900','£4,996,900'),
            ]
        },
        {
            'label':  'SCENARIO B — BASE  (Recommended)',
            'desc':   'Founders on lean but liveable UK salary. Full 8-person Pakistan team as modelled in '
                      'Section 4.2. Matches the $251K raise with ~14–16 months of runway. This is the basis '
                      'for the investor deck projections and the recommended operating model.',
            'bg':     '2E7D3E',
            'accent': '39B54A',
            'row_bg': 'E8F5E9',
            'years': [
                ('Year 1', '£86,600',  '£58,776',  '£24,000', '£169,376', '£70,000',  '(£99,376)'),
                ('Year 2', '£136,300', '£71,416',  '£43,100', '£250,816', '£337,140', '£86,324'),
                ('Year 3', '£231,000', '£100,014', '£87,900', '£418,914', '£1,122,780','£703,866'),
                ('Year 4', '£284,000', '£119,448', '£166,800','£570,248', '£2,807,700','£2,237,452'),
                ('Year 5', '£337,000', '£137,934', '£325,900','£800,834', '£5,616,900','£4,816,066'),
            ]
        },
        {
            'label':  'SCENARIO C — FULL TEAM',
            'desc':   'Founders at near-market rates from Year 1. Full Pakistan team plus UK Sales hire from Year 2. '
                      'Additional PK hires accelerate product velocity. Requires either faster revenue ramp or '
                      'a larger raise ($350K+). Justified if Seed closes early.',
            'bg':     '1B4332',
            'accent': '39B54A',
            'row_bg': 'E8F5E9',
            'years': [
                ('Year 1', '£110,000', '£70,000',  '£28,000', '£208,000', '£70,000',  '(£138,000)'),
                ('Year 2', '£160,000', '£85,000',  '£50,000', '£295,000', '£337,140', '£42,140'),
                ('Year 3', '£260,000', '£120,000', '£100,000','£480,000', '£1,122,780','£642,780'),
                ('Year 4', '£320,000', '£145,000', '£185,000','£650,000', '£2,807,700','£2,157,700'),
                ('Year 5', '£400,000', '£165,000', '£370,000','£935,000', '£5,616,900','£4,681,900'),
            ]
        },
    ]

    hdrs_sc = ['Year', 'UK Team', 'PK Team', 'Infra / Tools', 'Total OpEx', 'Gross Profit', 'Net P&L']

    for sc in scenarios:
        # Scenario label bar
        tbl_lbl = doc.add_table(rows=1, cols=1)
        tbl_lbl.columns[0].width = Inches(7.0)
        c_lbl = tbl_lbl.rows[0].cells[0]
        set_cell_bg(c_lbl, sc['bg'])
        set_cell_borders(c_lbl, colour=sc['accent'])
        set_cell_padding(c_lbl, 100, 100, 120, 120)
        cell_para(c_lbl, sc['label'], bold=True, size=11, colour=GREEN)

        add_para(doc, sc['desc'], size=9, colour=DARK_GREY, space_before=4, space_after=5)

        tbl_sc = doc.add_table(rows=1 + len(sc['years']), cols=7)
        tbl_sc.alignment = WD_TABLE_ALIGNMENT.CENTER
        w_sc = [Inches(0.6), Inches(0.98), Inches(0.98), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.1)]
        for i, w in enumerate(w_sc):
            for row in tbl_sc.rows:
                row.cells[i].width = w
        for i, h in enumerate(hdrs_sc):
            c = tbl_sc.rows[0].cells[i]
            set_cell_bg(c, sc['bg']); set_cell_borders(c, colour=sc['accent'])
            set_cell_padding(c, 75, 75, 75, 75)
            cell_para(c, h, bold=True, size=8.5, colour=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        for ri, r_ in enumerate(sc['years']):
            shade = ri % 2 == 0
            bg = sc['row_bg'] if shade else 'FFFFFF'
            for ci, val in enumerate(r_):
                c = tbl_sc.rows[ri + 1].cells[ci]
                if ci == 6:
                    is_neg = val.startswith('(')
                    set_cell_bg(c, 'FEE2E2' if is_neg else 'D1FAE5')
                else:
                    set_cell_bg(c, bg)
                set_cell_borders(c, colour='C3E6CB')
                set_cell_padding(c, 65, 65, 75, 75)
                cell_para(c, val, bold=(ci in (0, 6)), size=8.5,
                          align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_para(doc,
        'Gross profit values: Y1 = £70K · Y2 = £337K · Y3 = £1.12M · Y4 = £2.81M · Y5 = £5.62M '
        '(ARR × gross margin % converted at 0.79 USD/GBP). All Net P&L figures are pre-tax EBITDA proxies.',
        size=8, colour=MID_GREY, italic=True, space_before=2, space_after=4)

    add_divider(doc)
    doc.add_page_break()


# ── Section 6 – Consolidated P&L ─────────────────────────────────────────────

def build_pl(doc):
    section_banner(doc, '06   CONSOLIDATED PROFIT & LOSS — BASE SCENARIO',
                   'All figures in GBP · Pre-tax EBITDA proxy · USD/GBP 0.79')

    add_para(doc,
        'The consolidated P&L below tracks the business from pre-revenue through profitable scale under '
        'the Base Scenario (Scenario B). COGS is the direct AI + compute cost only. Gross Profit is '
        'after COGS. Net P&L deducts full operating expenditure. Year 2 is the first profitable year.',
        size=9.5, colour=DARK_GREY, space_before=2, space_after=6)

    hdrs = ['Line Item', 'Year 1\n2026–27', 'Year 2\n2027–28', 'Year 3\n2028–29',
            'Year 4\n2029–30', 'Year 5\n2030–31']

    pl_rows = [
        # (values, bg_override or None)
        (('ARR (USD)',             '$108,000',  '$540,000',    '$1,800,000',  '$4,500,000',  '$9,000,000'),  None),
        (('ARR (GBP @ 0.79)',      '£85,320',   '£426,600',    '£1,422,000',  '£3,555,000',  '£7,110,000'),  None),
        (('COGS — Direct (GBP)',   '(£14,751)', '(£38,394)',   '(£99,540)',   '(£213,300)',  '(£355,500)'),  'FEE2E2'),
        (('Gross Profit (GBP)',    '£70,569',   '£388,206',    '£1,322,460',  '£3,341,700',  '£6,754,500'),  'D1FAE5'),
        (('Gross Margin %',        '82.7%',     '91.0%',       '93.0%',       '94.0%',       '95.0%'),       'D1FAE5'),
        (('',)*6,                                                                                             'FFFFFF'),
        (('UK Team Salaries',      '(£86,600)', '(£136,300)',  '(£231,000)',  '(£284,000)',  '(£337,000)'),  None),
        (('Pakistan Team Cost',    '(£58,776)', '(£71,416)',   '(£100,014)',  '(£119,448)',  '(£137,934)'),  None),
        (('Infrastructure & Tools','(£24,000)', '(£43,100)',   '(£87,900)',   '(£166,800)',  '(£325,900)'),  None),
        (('Total OpEx',            '(£169,376)','(£250,816)',  '(£418,914)',  '(£570,248)',  '(£800,834)'),  'FEE2E2'),
        (('',)*6,                                                                                             'FFFFFF'),
        (('Net P&L (EBITDA proxy)','(£98,807)', '£137,390',    '£903,546',    '£2,771,452',  '£5,953,666'),  None),
        (('Cumulative P&L',        '(£98,807)', '£38,583',     '£942,129',    '£3,713,581',  '£9,667,247'),  None),
    ]

    tbl = doc.add_table(rows=1 + len(pl_rows), cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(2.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0)]
    for i, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[i].width = w
    for i, h in enumerate(hdrs):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, '1B4332'); set_cell_borders(c, '1B4332')
        set_cell_padding(c, 80, 80, 80, 80)
        cell_para(c, h, bold=True, size=8.5, colour=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    for ri, (values, override_bg) in enumerate(pl_rows):
        is_net   = ri == 11
        is_cumul = ri == 12
        shade    = ri % 2 == 0
        default_bg = hex_from_rgb(LIGHT_GREEN) if shade else 'FFFFFF'
        bg = override_bg if override_bg else default_bg
        if is_net or is_cumul:
            bg = '1B4332'

        for ci, val in enumerate(values):
            c = tbl.rows[ri + 1].cells[ci]
            # Net P&L — colour by value
            if (is_net or is_cumul) and ci > 0:
                cell_bg = '1B4332'
            elif (is_net or is_cumul) and ci == 0:
                cell_bg = '1B4332'
            else:
                cell_bg = bg
            set_cell_bg(c, cell_bg)
            set_cell_borders(c, colour='C3E6CB' if not (is_net or is_cumul) else '1B4332')
            set_cell_padding(c, 70, 70, 80, 80)
            text_col = WHITE if (is_net or is_cumul) else DARK_GREY
            if (is_net or is_cumul) and ci > 0 and val and not val.startswith('('):
                text_col = GREEN
            al = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            cell_para(c, str(val), bold=(is_net or is_cumul or ri in (3, 4, 9)),
                      size=8.5, colour=text_col, align=al)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Callout boxes
    callouts = [
        ('Year 2', 'Net P&L turns positive — first profitable year', '2E7D3E', GREEN),
        ('Year 2 End', 'Cumulative P&L positive — raise fully returned', '2E7D3E', GREEN),
        ('Year 3', '£900K+ net profit — self-funding further expansion', '1B4332', WHITE),
        ('Year 5', '£9.6M cumulative — strong Series A justification', '1B4332', WHITE),
    ]

    tbl_co = doc.add_table(rows=1, cols=4)
    tbl_co.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (yr, msg, bg, tc) in enumerate(callouts):
        c = tbl_co.rows[0].cells[i]
        c.width = Inches(1.75)
        set_cell_bg(c, bg); set_cell_borders(c, colour='39B54A')
        set_cell_padding(c, 100, 100, 80, 80)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(yr + '\n')
        r1.bold = True; r1.font.size = Pt(9); r1.font.color.rgb = GREEN; r1.font.name = 'Calibri'
        r2 = p.add_run(msg)
        r2.font.size = Pt(7.5); r2.font.color.rgb = tc; r2.font.name = 'Calibri'

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_divider(doc)
    doc.add_page_break()


# ── Section 7 – Raise Deployment ─────────────────────────────────────────────

def build_raise(doc):
    section_banner(doc, '07   RAISE DEPLOYMENT & RUNWAY',
                   '$251,000 Pre-Seed SAFE — capital allocation and runway by scenario')

    add_para(doc,
        'The $251,000 Pre-Seed raise (~£198,290 at 0.79) funds 12 declared months of runway to hit '
        'Seed-qualifying milestones: 150 paying clients and $10K MRR by Month 12, staging MVP live '
        'by Month 3. Under the Base Scenario, the same capital extends to 14–16 months. Under the '
        'Lean Scenario, it reaches 18–20 months.',
        size=9.5, colour=DARK_GREY, space_before=2, space_after=6)

    add_para(doc, '7.1  Capital Allocation', bold=True, size=11,
             colour=DARK_TEAL, space_before=2, space_after=4)

    hdrs = ['Category', 'USD', 'GBP (~0.79)', '% of Raise', 'Notes']
    alloc_rows = [
        ('UK Team Salaries (Y1 lean)',               '$110,000', '£86,900', '43.8%', 'Founders + BD; Employer NI included'),
        ('Pakistan Engineering Team (Y1)',           '$65,000',  '£51,350', '25.9%', 'Backend, frontend, DevOps, security'),
        ('Marketing & Sales',                        '$20,000',  '£15,800', '8.0%',  'SEO, LinkedIn ads, partner outreach'),
        ('AWS Infrastructure (net of credits)',      '$12,000',  '£9,480',  '4.8%',  'Staging + early prod environment'),
        ('Legal & Compliance',                       '$12,000',  '£9,480',  '4.8%',  'SAFE docs, GDPR, contracts, insurance'),
        ('Tools, APIs & SaaS Subscriptions',         '$8,000',   '£6,320',  '3.2%',  'GitHub, security APIs, Notion, Slack'),
        ('Contingency & Working Capital',            '$24,000',  '£18,960', '9.6%',  '~3 months minimum burn buffer'),
    ]
    widths = [Inches(2.1), Inches(0.9), Inches(0.9), Inches(0.85), Inches(2.0)]
    tbl = make_table(doc, hdrs, alloc_rows, widths, font_size=8.5)
    # Add total row
    total_row(tbl, len(alloc_rows) + 1 if False else -1, [])  # skip — add manually below

    # Manual total row appended
    tbl2 = doc.add_table(rows=1, cols=5)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (w, val) in enumerate(zip(widths, ['TOTAL', '$251,000', '£198,290', '100%', ''])):
        c = tbl2.rows[0].cells[i]
        c.width = w
        set_cell_bg(c, '1B4332'); set_cell_borders(c, '1B4332')
        set_cell_padding(c, 80, 80, 80, 80)
        al = WD_ALIGN_PARAGRAPH.LEFT if i in (0, 4) else WD_ALIGN_PARAGRAPH.CENTER
        cell_para(c, val, bold=True, size=8.5, colour=WHITE, align=al)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_para(doc, '7.2  Runway Comparison by Scenario', bold=True, size=11,
             colour=DARK_TEAL, space_before=4, space_after=4)

    hdrs2 = ['Scenario', 'Monthly Burn', 'Annual Burn (Y1)', 'Runway on $251K', 'Break-even Month']
    rows2 = [
        (('A — Lean',      '~£9,500–£11,000',  '~£130,000', '18–20 months', 'Month ~22'), 'E8F5E9'),
        (('B — Base',      '~£13,500–£15,000', '~£169,000', '14–16 months', 'Month ~28'), 'D1FAE5'),
        (('C — Full Team', '~£16,500–£18,500', '~£208,000', '11–12 months', 'Month ~32'), 'E8F5E9'),
    ]
    widths2 = [Inches(1.4), Inches(1.4), Inches(1.3), Inches(1.4), Inches(1.4)]
    make_table(doc, hdrs2, rows2, widths2, font_size=9)

    add_para(doc,
        'Monthly burn = total OpEx ÷ 12. Revenue ramp reduces net cash consumption from Month 6 onwards. '
        'The Seed close at Month 15 ($750K target) refills the runway before the Base Scenario reaches '
        'zero — no bridge financing required.',
        size=8, colour=MID_GREY, italic=True, space_before=5, space_after=4)

    add_divider(doc)
    doc.add_page_break()


# ── Section 8 – Assumptions & Sensitivity ────────────────────────────────────

def build_assumptions(doc):
    section_banner(doc, '08   KEY ASSUMPTIONS & SENSITIVITY',
                   'Model inputs and what changes if they shift')

    add_para(doc, '8.1  Core Model Assumptions', bold=True, size=11,
             colour=DARK_TEAL, space_before=2, space_after=4)

    assumptions = [
        ('USD / GBP exchange rate',     '0.79',
         '5% GBP weakening adds ~£4K to PK costs in Y1 — low impact'),
        ('Claude API pricing',          '$3.00/1M input · $15.00/1M output · $0.30/1M cached',
         'Prompt caching cuts effective input cost ~90%'),
        ('Pakistan salary escalation',  '10–15% per year',
         'Competitive market; senior roles escalate faster'),
        ('ARPU at maturity',            '$150 / month',
         'Mix of $199 SOC Pro (65%) + $499 Enterprise (20%)'),
        ('Free-to-paid conversion',     '8–12%',
         'Freemium benchmark; SOCVault target 10%'),
        ('Monthly churn',               '2.5% Y1 → 1.2% Y3',
         'SMB churn is high early; compliance stickiness reduces it'),
        ('AWS Activate credits',        '$25,000–$100,000',
         'Applied for but not confirmed — contingency covers if not received'),
        ('Seed round close',            'Month 15 (Aug 2027)',
         '$750K SAFE or priced round at $5–8M cap'),
        ('UK Employer NI rate',         '13.8%',
         'Applied to all employed UK staff only'),
        ('VAT treatment',               'VAT registered from first invoice',
         'Output VAT collected; input VAT recovered on costs'),
    ]

    hdrs = ['Assumption', 'Value Used', 'Sensitivity / Comment']
    widths = [Inches(2.0), Inches(2.2), Inches(2.8)]
    make_table(doc, hdrs, assumptions, widths, font_size=8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_para(doc, '8.2  Impact Sensitivity', bold=True, size=11,
             colour=DARK_TEAL, space_before=4, space_after=4)

    levels = [
        ('HIGH IMPACT', 'EF4444', [
            'ARPU shift: +$20 ARPU (more Enterprise conversions) adds ~£190K to Y3 ARR alone',
            'Churn rate: dropping from 2.5% to 1.5% in Y1 retains 10 more clients — £21,600 ARR',
            'Pakistan salary escalation: 20%/yr vs 10%/yr adds ~£14K to Y3 OpEx — manageable',
            'AWS credits not received: adds £8,000–£18,000 to Y1 infra cost — within contingency buffer',
        ]),
        ('MEDIUM IMPACT', 'F59E0B', [
            'Founder salary draw: each additional £10K drawn adds £11,380 (with NI) to annual OpEx',
            'Free-to-paid conversion: +1% improvement generates ~15 additional clients per year',
            'Claude API price change: 20% reduction saves ~£3K in Y1 — small relative to OpEx',
        ]),
        ('LOW IMPACT', '39B54A', [
            'Exchange rate: 5% GBP move changes PK cost by ~£3,000–£7,000/year',
            'SaaS tool costs: <2% of OpEx at all stages — not worth optimising early',
        ]),
    ]

    for level, colour, items in levels:
        tbl_lv = doc.add_table(rows=1, cols=1)
        tbl_lv.columns[0].width = Inches(7.0)
        c_lv = tbl_lv.rows[0].cells[0]
        set_cell_bg(c_lv, colour); set_cell_borders(c_lv, colour=colour)
        set_cell_padding(c_lv, 80, 80, 100, 100)
        cell_para(c_lv, level, bold=True, size=9, colour=WHITE)

        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(item)
            r.font.size = Pt(9)
            r.font.name = 'Calibri'
            r.font.color.rgb = DARK_GREY
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_divider(doc)
    doc.add_page_break()


# ── Section 9 – Disclaimer ────────────────────────────────────────────────────

def build_disclaimer(doc):
    section_banner(doc, '09   IMPORTANT NOTICES & DISCLAIMER')

    add_para(doc,
        'This document is a working financial model prepared by SOCVault Ltd for internal planning '
        'purposes and discussion with prospective investors and strategic partners. It contains '
        'forward-looking statements and projections based on assumptions about future events and '
        'conditions that may not occur. Actual results may differ materially from those projected.',
        size=9, colour=DARK_GREY, space_before=2, space_after=6)

    notices = [
        'All revenue figures are projections, not guarantees. SOCVault is pre-revenue at the date of this document.',
        'Pakistan contractor rates are based on June 2026 market data and are subject to change.',
        'Claude API pricing is based on published Anthropic rates as of June 2026 and is subject to change.',
        'GBP/USD exchange rate of 0.79 is indicative; currency fluctuation creates material risk for a UK company with USD-denominated costs.',
        'AWS Activate credits are applied for but not confirmed at the time of this document.',
        'This document does not constitute a financial promotion under the Financial Services and Markets Act 2000 (as amended).',
        'Recipients should conduct their own due diligence and obtain independent financial and legal advice before making any investment decision.',
        'SOCVault Ltd is registered in England & Wales. SOCVault (SMC-PVT) Ltd is registered in Pakistan (Reg. BG96012).',
    ]

    for notice in notices:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(notice)
        r.font.size = Pt(8.5)
        r.font.name = 'Calibri'
        r.font.color.rgb = DARK_GREY

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Final stamp-style footer
    if os.path.exists(LOGO_ICON_PATH):
        p_icon = doc.add_paragraph()
        p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_icon.paragraph_format.space_before = Pt(6)
        p_icon.paragraph_format.space_after  = Pt(4)
        r_ic = p_icon.add_run()
        r_ic.add_picture(LOGO_ICON_PATH, width=Inches(0.55))

    p_final = doc.add_paragraph()
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_final.add_run(
        'SOCVault Ltd  ·  Registered in England & Wales\n'
        'SOCVault (SMC-PVT) Ltd  ·  Reg. BG96012  ·  Karachi, Pakistan\n'
        'socvault.io  ·  CONFIDENTIAL'
    )
    r_f.font.size = Pt(8)
    r_f.font.color.rgb = MID_GREY
    r_f.font.name = 'Calibri'


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(9.5)

    setup_header_footer(doc)
    build_cover(doc)
    build_exec_summary(doc)
    build_unit_economics(doc)
    build_revenue(doc)
    build_opex(doc)
    build_scenarios(doc)
    build_pl(doc)
    build_raise(doc)
    build_assumptions(doc)
    build_disclaimer(doc)

    out = 'BusinessProposal/SOCVault-Revenue-OpEx-Model.docx'
    doc.save(out)
    print(f'Saved → {out}')


if __name__ == '__main__':
    main()
